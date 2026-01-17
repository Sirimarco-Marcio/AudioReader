from pathlib import Path
from pypdf import PdfReader
from docx import Document
from app.core.vision import VisionAI

def extrair_conteudo(caminho_arquivo, usar_vision=False, callback_progresso=None):
    """
    Lê o arquivo e retorna o texto completo.
    callback_progresso: função para atualizar barra de progresso (opcional).
    """
    caminho = Path(caminho_arquivo)
    extensao = caminho.suffix.lower()
    texto_final = ""
    
    # Prepara a IA se solicitado
    vision_service = VisionAI()
    if usar_vision:
        sucesso = vision_service.inicializar()
        if not sucesso:
            usar_vision = False # Desativa se falhar no load

    if extensao == '.pdf':
        try:
            reader = PdfReader(caminho)
            total_paginas = len(reader.pages)
            
            for i, page in enumerate(reader.pages):
                # 1. Extrai Texto
                texto_pag = page.extract_text() or ""
                
                # 2. Extrai e Descreve Imagens (Se ativado)
                descricoes_visuais = []
                if usar_vision:
                    # pypdf extrai imagens como objetos binários
                    for image_file_object in page.images:
                        descricao = vision_service.descrever_imagem(image_file_object.data)
                        if descricao:
                            print(f"   👁️  Imagem detectada na pág {i+1}: {descricao}")
                            descricoes_visuais.append(descricao)
                
                # 3. Montagem da Página
                # Estratégia: Ler o texto primeiro, depois descrever as imagens daquela página
                texto_final += texto_pag + "\n"
                if descricoes_visuais:
                    texto_final += "\n--- CONTEÚDO VISUAL DETECTADO NA PÁGINA ---\n"
                    texto_final += "\n".join(descricoes_visuais)
                    texto_final += "\n-------------------------------------------\n"
                
                # Atualiza progresso
                if callback_progresso:
                    callback_progresso(i + 1, total_paginas, "Lendo PDF/Imagens")
                    
        except Exception as e:
            print(f"Erro crítico no PDF: {e}")
            return None

    elif extensao in ['.docx', '.doc']:
        doc = Document(caminho)
        total = len(doc.paragraphs)
        for i, para in enumerate(doc.paragraphs):
            texto_final += para.text + "\n"
            if callback_progresso and i % 10 == 0:
                callback_progresso(i + 1, total, "Lendo DOCX")

    return texto_final