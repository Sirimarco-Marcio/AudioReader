import os
import sys
import time
import shutil
from pathlib import Path
from pypdf import PdfReader
from docx import Document
import asyncio
import edge_tts

from gtts import gTTS

# Cores para o terminal
VERDE = '\033[92m'
AMARELO = '\033[93m'
AZUL = '\033[94m'
RESET = '\033[0m'

def limpar_tela():
    os.system('cls' if os.name == 'nt' else 'clear')

def barra_progresso(iteracao, total, prefixo='', tamanho=30):
    if total == 0: return
    percentual = ("{0:.1f}").format(100 * (iteracao / float(total)))
    preenchido = int(tamanho * iteracao // total)
    barra = '█' * preenchido + '-' * (tamanho - preenchido)
    sys.stdout.write(f'\r{prefixo} |{barra}| {percentual}% Completo')
    sys.stdout.flush()

def listar_arquivos_recursivo():
    extensoes = {'.pdf', '.docx', '.doc'}
    arquivos_encontrados = []
    
    # Percorre a pasta atual e todas as subpastas
    for root, dirs, files in os.walk('.'):
        # Pula as pastas de áudio já geradas para não duplicar no menu
        if "AUDIO -" in root:
            continue
            
        for file in files:
            caminho_completo = Path(root) / file
            # Verifica extensão e ignora arquivos temporários do Word (~$)
            if caminho_completo.suffix.lower() in extensoes and not file.startswith('~$'):
                # Salva o caminho relativo (ex: "PDFs/arquivo.pdf")
                arquivos_encontrados.append(str(caminho_completo))
                
    return sorted(arquivos_encontrados)

def criar_pasta_destino(arquivo_path):
    # arquivo_path aqui pode ser "PDFs/arquivo.pdf"
    # Pegamos apenas o nome do arquivo sem a pasta pai para criar o output
    nome_arquivo = arquivo_path.stem
    nome_pasta = f"AUDIO - {nome_arquivo}"
    
    if not os.path.exists(nome_pasta):
        os.makedirs(nome_pasta)
        print(f"\n{VERDE}📂 Pasta criada na raiz: {nome_pasta}{RESET}")
    
    return Path(nome_pasta)

def extrair_texto(caminho):
    texto_completo = ""
    caminho_obj = Path(caminho)
    extensao = caminho_obj.suffix.lower()

    try:
        if extensao == '.pdf':
            reader = PdfReader(caminho)
            total = len(reader.pages)
            print(f"{AZUL}📖 Lendo PDF ({total} páginas)...{RESET}")
            
            for i, page in enumerate(reader.pages):
                texto = page.extract_text()
                if texto:
                    texto_completo += texto + " "
                barra_progresso(i + 1, total, prefixo='Extraindo Texto')
                
        elif extensao in ['.docx', '.doc']:
            doc = Document(caminho)
            total = len(doc.paragraphs)
            print(f"{AZUL}📝 Lendo DOCX ({total} parágrafos)...{RESET}")
            
            for i, para in enumerate(doc.paragraphs):
                texto_completo += para.text + " "
                if i % 10 == 0 or i == total - 1:
                    barra_progresso(i + 1, total, prefixo='Extraindo Texto')
        
        print() 
        return texto_completo

    except Exception as e:
        print(f"\n{AMARELO}❌ Erro ao ler arquivo: {e}{RESET}")
        return None

def gerar_audio(texto, caminho_saida):
    if not texto or not texto.strip():
        print(f"{AMARELO}⚠️  O arquivo parece estar vazio ou é uma imagem.{RESET}")
        return

    print(f"{AZUL}🔊 Gerando Áudio com IA (Isso pode levar um tempo, aguarde...){RESET}")
    
    try:
        inicio = time.time()
        tts = gTTS(text=texto, lang='pt', slow=False)
        tts.save(caminho_saida)
        fim = time.time()
        print(f"{VERDE}✅ Áudio concluído em {fim - inicio:.2f} segundos!{RESET}")
    except Exception as e:
        print(f"{AMARELO}❌ Erro na conexão com o Google TTS: {e}{RESET}")

# --- NOVA FUNÇÃO GERAR_AUDIO (Substitua a antiga inteira por esta) ---
def gerar_audio(texto, caminho_saida):
    if not texto or not texto.strip():
        print(f"{AMARELO}⚠️  O arquivo parece estar vazio ou é uma imagem.{RESET}")
        return

    print(f"{AZUL}🔊 Gerando Áudio Neural (Isso tem qualidade de podcast!)...{RESET}")
    
    # Vozes disponíveis em PT-BR:
    # "pt-BR-FranciscaNeural" (Feminina - Muito natural)
    # "pt-BR-AntonioNeural"   (Masculina - Ótima para textos técnicos)
    VOZ = "pt-BR-AntonioNeural" 
    
    async def _gerar():
        communicate = edge_tts.Communicate(texto, VOZ)
        await communicate.save(caminho_saida)

    try:
        inicio = time.time()
        # O edge-tts é assíncrono, então precisamos desse loop para rodar no script normal
        asyncio.run(_gerar())
        fim = time.time()
        print(f"{VERDE}✅ Áudio concluído em {fim - inicio:.2f} segundos!{RESET}")
    except Exception as e:
        print(f"{AMARELO}❌ Erro na geração neural: {e}{RESET}")
        print("Verifique sua conexão com a internet.")
        
def main():
    while True:
        limpar_tela()
        print(f"{VERDE}=== CONVERSOR DE DOCUMENTOS PARA AUDIOBOOK (V2) ==={RESET}")
        print("Procurando arquivos nas pastas DOCs, PDFs e Raiz...\n")

        arquivos = listar_arquivos_recursivo()

        if not arquivos:
            print(f"{AMARELO}Nenhum arquivo encontrado nas pastas.{RESET}")
            print(f"Certifique-se que os arquivos estão em 'PDFs', 'DOCs' ou na raiz.")
            input("\nPressione Enter para tentar novamente...")
            continue

        # Menu Dinâmico
        for i, arq in enumerate(arquivos):
            # Mostra o caminho (ex: PDFs/arquivo.pdf) para você saber de onde veio
            print(f"[{i+1}] {arq}")
        
        print(f"\n[0] Sair")

        opcao = input(f"\n{AZUL}Escolha o número do arquivo: {RESET}")

        if opcao == '0':
            print("Saindo...")
            break

        try:
            indice = int(opcao) - 1
            if 0 <= indice < len(arquivos):
                caminho_str = arquivos[indice]
                arquivo_escolhido = Path(caminho_str)
                
                print(f"\nProcessando: {arquivo_escolhido}")

                # 1. Estrutura de Pastas (Cria na raiz do projeto)
                pasta_destino = criar_pasta_destino(arquivo_escolhido)
                
                # 2. Copia o original para a pasta (Backup)
                shutil.copy2(arquivo_escolhido, pasta_destino / arquivo_escolhido.name)
                print(f"📄 Cópia do documento salva.")

                # 3. Extração
                texto = extrair_texto(arquivo_escolhido)

                # 4. Geração do Áudio
                if texto:
                    nome_audio = f"AUDIO - {arquivo_escolhido.stem}.mp3"
                    caminho_audio = pasta_destino / nome_audio
                    gerar_audio(texto, caminho_audio)

                input(f"\n{VERDE}Pressione Enter para voltar ao menu...{RESET}")
            else:
                print(f"{AMARELO}Opção inválida!{RESET}")
                time.sleep(1)
        except ValueError:
            print(f"{AMARELO}Digite apenas números.{RESET}")
            time.sleep(1)

if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\nOperação cancelada.")